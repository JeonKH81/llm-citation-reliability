#!/usr/bin/env python3
"""Build standalone Word tables for the main manuscript."""
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "main_tables_JMIR.docx"


TABLE1 = {
    "caption": "Table 1. Characteristics of review topics and generated references",
    "headers": [
        "Publication-volume level",
        "Topic",
        "Topic type",
        "PubMed-indexed articles",
        "Reviews generated",
        "References available for verification",
        "References with stated identifier",
    ],
    "rows": [
        ["Low", "Transcaval access", "Procedure", "191", "30", "900", "891"],
        ["Low", "Cardiac contractility modulation", "Device therapy", "274", "30", "900", "897"],
        ["Low", "Coronary slow flow phenomenon", "Disease or phenomenon", "477", "30", "902", "902"],
        ["Moderate", "Subcutaneous implantable cardioverter defibrillator", "Device", "1,175", "30", "900", "900"],
        ["Moderate", "Bioresorbable vascular scaffold", "Device or stent technology", "1,389", "30", "900", "900"],
        ["Moderate", "Peripartum cardiomyopathy", "Disease", "1,852", "30", "900", "899"],
        ["High", "Sacubitril and valsartan", "Drug therapy", "3,187", "30", "890", "889"],
        ["High", "Left atrial appendage closure", "Procedure", "3,780", "30", "868", "868"],
        ["High", "Cardiac amyloidosis", "Disease", "4,849", "30", "890", "890"],
    ],
    "widths": [1.15, 2.20, 1.45, 1.15, 1.05, 1.35, 1.30],
}


TABLE2 = {
    "caption": "Table 2. Reference verification outcomes by model",
    "headers": ["Verification outcome", "GPT-5.5", "Claude Opus 4.8", "Gemini 3.5 Flash"],
    "rows": [
        ["Verified", "2330", "1796", "1865"],
        ["Verified benign", "59", "42", "37"],
        ["Real but not in PubMed", "0", "28", "0"],
        ["Field error", "157", "164", "103"],
        ["Misattribution", "154", "597", "692"],
        ["Fabrication", "0", "7", "5"],
        ["Omission", "0", "14", "0"],
    ],
    "widths": [3.30, 1.60, 1.90, 1.90],
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_run(run, bold=False, size=10):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    format_run(r, bold=True, size=10)


def add_word_table(doc, spec):
    add_caption(doc, spec["caption"])
    table = doc.add_table(rows=1, cols=len(spec["headers"]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for j, header in enumerate(spec["headers"]):
        cell = hdr.cells[j]
        set_cell_shading(cell, "F2F2F2")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        format_run(run, bold=True, size=9)
    for row_data in spec["rows"]:
        row = table.add_row()
        for j, value in enumerate(row_data):
            cell = row.cells[j]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            format_run(run, size=9)
    set_table_width(table, spec["widths"])
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("Main manuscript tables")
    format_run(r, bold=True, size=12)

    add_word_table(doc, TABLE1)
    add_word_table(doc, TABLE2)

    doc.save(OUT)
    print(f"saved {OUT}")


if __name__ == "__main__":
    main()
